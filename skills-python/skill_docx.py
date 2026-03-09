# Required pip packages: python-docx
import sys
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_wine_list_docx(wines):
    doc = Document()
    
    # Add title
    title = doc.add_heading('Wine List', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add introduction
    doc.add_paragraph('A curated selection of fine wines for your enjoyment.')
    doc.add_paragraph('')  # Add empty line
    
    # Add wine entries
    for wine in wines:
        # Add wine heading
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Wine name in bold
        run_name = p.add_run(wine['name'])
        run_name.bold = True
        run_name.font.size = Pt(12)
        
        # Add details as normal text
        details = f" - Region: {wine['region']}, Variety: {wine['variety']}, Year: {wine['year']}"
        p.add_run(details)
        
        # Add price
        price_run = p.add_run(f" | Price: ${wine['price']:.2f}")
        price_run.italic = True
        
        # Add empty line after each wine
        doc.add_paragraph('')
    
    # Save document
    doc.save('wine_list.docx')
    print("Wine list generated successfully as 'wine_list.docx'")

def main():
    if len(sys.argv) > 1:
        # Read from command line argument
        task = sys.argv[1]
    else:
        # Read from stdin
        task = sys.stdin.read().strip()
    
    if not task:
        print("Please provide wine data in JSON format or as a list of dictionaries.")
        return
    
    try:
        import json
        wines = json.loads(task)
        if isinstance(wines, list):
            generate_wine_list_docx(wines)
        else:
            print("Input must be a JSON array of wine objects.")
    except Exception as e:
        print(f"Error processing input: {e}")
        print("Please provide wine data as a JSON array of objects with fields: name, region, variety, year, price")

if __name__ == "__main__":
    main()